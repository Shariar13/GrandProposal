"""
Export Engine for Proposals
Exports to DOCX, PDF, LaTeX, and other formats with professional formatting
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from io import BytesIO
import re
from typing import Dict, List, Any
import markdown
from datetime import datetime


class ProposalExportEngine:
    """
    Professional export engine for grant proposals
    Supports: DOCX, PDF, LaTeX, Markdown, HTML
    """

    def __init__(self):
        self.page_size = A4
        self.margin = 1 * inch

    def export_to_docx(
        self,
        proposal_data: Dict[str, Any],
        include_citations: bool = True,
        include_cover_page: bool = True
    ) -> BytesIO:
        """
        Export proposal to professional DOCX format
        """
        doc = Document()

        # Set up styles
        self._setup_docx_styles(doc)

        # Cover page
        if include_cover_page:
            self._add_docx_cover_page(doc, proposal_data)
            doc.add_page_break()

        # Table of Contents placeholder
        doc.add_heading('Table of Contents', level=1)
        doc.add_paragraph('(Table of contents will be auto-generated in Word)')
        doc.add_page_break()

        # Add each section
        sections = proposal_data.get('sections', [])
        for section in sections:
            # Section heading
            doc.add_heading(section['name'], level=1)

            # Section content
            content = section.get('content', '')
            self._add_formatted_content(doc, content)

            # Add page break after major sections
            if section.get('page_break_after', False):
                doc.add_page_break()

        # References/Bibliography
        if include_citations:
            doc.add_page_break()
            doc.add_heading('References', level=1)
            self._add_docx_bibliography(doc, proposal_data)

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    def _setup_docx_styles(self, doc: Document):
        """Setup professional styles for DOCX"""
        styles = doc.styles

        # Heading 1 style
        h1 = styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(31, 78, 121)

        # Heading 2 style
        h2 = styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(79, 129, 189)

        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)

    def _add_docx_cover_page(self, doc: Document, proposal_data: Dict):
        """Add professional cover page"""
        # Title
        title = doc.add_heading(proposal_data.get('title', 'Grant Proposal'), level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph('')  # Spacer

        # Proposal type
        type_para = doc.add_paragraph()
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_para.add_run(proposal_data.get('proposal_type', ''))
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph('')  # Spacer

        # Metadata table
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        metadata = [
            ('Principal Investigator:', proposal_data.get('pi_name', '')),
            ('Institution:', proposal_data.get('institution', '')),
            ('Date:', datetime.now().strftime('%B %d, %Y')),
            ('Keywords:', proposal_data.get('keywords', '')),
            ('Funding Amount:', proposal_data.get('funding_amount', ''))
        ]

        for idx, (label, value) in enumerate(metadata):
            table.rows[idx].cells[0].text = label
            table.rows[idx].cells[1].text = value

    def _add_formatted_content(self, doc: Document, content: str):
        """Add formatted content with markdown support"""
        # Split into paragraphs
        paragraphs = content.split('\n\n')

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check for markdown headings
            if para.startswith('###'):
                doc.add_heading(para.replace('###', '').strip(), level=3)
            elif para.startswith('##'):
                doc.add_heading(para.replace('##', '').strip(), level=2)
            elif para.startswith('#'):
                doc.add_heading(para.replace('#', '').strip(), level=1)
            # Check for bullet points
            elif para.startswith('- ') or para.startswith('* '):
                lines = para.split('\n')
                for line in lines:
                    if line.strip():
                        doc.add_paragraph(
                            line.lstrip('- *').strip(),
                            style='List Bullet'
                        )
            # Regular paragraph
            else:
                p = doc.add_paragraph(para)
                # Add citations formatting
                p.text = self._format_citations_in_text(p.text)

    def _format_citations_in_text(self, text: str) -> str:
        """Format citations in text"""
        # Find patterns like (Author, Year) and make them italic
        citation_pattern = r'\(([A-Z][a-z]+(?:\s+et\s+al\.)?),?\s+\d{4}\)'
        # This would need actual citation formatting
        return text

    def _add_docx_bibliography(self, doc: Document, proposal_data: Dict):
        """Add bibliography in APA format"""
        citations = proposal_data.get('bibliography', {}).get('apa', [])

        for citation in citations:
            p = doc.add_paragraph(citation, style='Normal')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    def export_to_pdf(
        self,
        proposal_data: Dict[str, Any],
        include_citations: bool = True
    ) -> BytesIO:
        """
        Export proposal to professional PDF format
        """
        buffer = BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        # Container for PDF elements
        elements = []

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4e79'),
            spaceAfter=30,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            textColor=colors.HexColor('#1f4e79'),
            spaceAfter=12,
            spaceBefore=12,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            leading=16,
            spaceAfter=12,
            alignment=4  # Justify
        )

        # Title page
        elements.append(Spacer(1, 2*inch))
        elements.append(Paragraph(proposal_data.get('title', 'Grant Proposal'), title_style))
        elements.append(Spacer(1, 0.5*inch))
        elements.append(Paragraph(proposal_data.get('proposal_type', ''), styles['Heading2']))
        elements.append(Spacer(1, 0.3*inch))

        # Metadata table
        metadata_data = [
            ['PI:', proposal_data.get('pi_name', '')],
            ['Institution:', proposal_data.get('institution', '')],
            ['Date:', datetime.now().strftime('%B %d, %Y')],
            ['Keywords:', proposal_data.get('keywords', '')]
        ]

        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))

        elements.append(metadata_table)
        elements.append(PageBreak())

        # Add sections
        sections = proposal_data.get('sections', [])
        for idx, section in enumerate(sections):
            # Section heading
            elements.append(Paragraph(section['name'], heading_style))
            elements.append(Spacer(1, 0.1*inch))

            # Section content
            content = section.get('content', '')
            paragraphs = content.split('\n\n')

            for para in paragraphs:
                para = para.strip()
                if para:
                    # Remove markdown headings
                    para = re.sub(r'^#{1,6}\s*', '', para)
                    elements.append(Paragraph(para, body_style))

            elements.append(Spacer(1, 0.2*inch))

        # Bibliography
        if include_citations:
            elements.append(PageBreak())
            elements.append(Paragraph('References', heading_style))
            elements.append(Spacer(1, 0.2*inch))

            citations = proposal_data.get('bibliography', {}).get('apa', [])
            for citation in citations:
                elements.append(Paragraph(citation, body_style))

        # Build PDF
        doc.build(elements)
        buffer.seek(0)

        return buffer

    def export_to_latex(self, proposal_data: Dict[str, Any]) -> str:
        """
        Export proposal to LaTeX format
        """
        latex_parts = []

        # Document class and packages
        latex_parts.append(r'\documentclass[11pt,a4paper]{article}')
        latex_parts.append(r'\usepackage[utf8]{inputenc}')
        latex_parts.append(r'\usepackage{geometry}')
        latex_parts.append(r'\usepackage{times}')
        latex_parts.append(r'\usepackage{hyperref}')
        latex_parts.append(r'\usepackage{graphicx}')
        latex_parts.append(r'\usepackage{cite}')
        latex_parts.append(r'\geometry{margin=1in}')
        latex_parts.append(r'\setlength{\parindent}{0pt}')
        latex_parts.append(r'\setlength{\parskip}{1em}')
        latex_parts.append('')

        # Title and metadata
        latex_parts.append(r'\title{' + self._escape_latex(proposal_data.get('title', '')) + '}')
        latex_parts.append(r'\author{' + self._escape_latex(proposal_data.get('pi_name', '')) + '}')
        latex_parts.append(r'\date{\today}')
        latex_parts.append('')

        # Begin document
        latex_parts.append(r'\begin{document}')
        latex_parts.append(r'\maketitle')
        latex_parts.append(r'\tableofcontents')
        latex_parts.append(r'\newpage')
        latex_parts.append('')

        # Add sections
        sections = proposal_data.get('sections', [])
        for section in sections:
            latex_parts.append(r'\section{' + self._escape_latex(section['name']) + '}')
            latex_parts.append('')

            content = section.get('content', '')
            # Convert markdown to LaTeX
            content = self._markdown_to_latex(content)
            latex_parts.append(content)
            latex_parts.append('')

        # Bibliography
        latex_parts.append(r'\newpage')
        latex_parts.append(r'\section*{References}')
        citations = proposal_data.get('bibliography', {}).get('apa', [])
        latex_parts.append(r'\begin{enumerate}')
        for citation in citations:
            latex_parts.append(r'\item ' + self._escape_latex(citation))
        latex_parts.append(r'\end{enumerate}')

        # End document
        latex_parts.append(r'\end{document}')

        return '\n'.join(latex_parts)

    def _escape_latex(self, text: str) -> str:
        """Escape special LaTeX characters"""
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\^{}',
            '\\': r'\textbackslash{}',
        }

        for old, new in replacements.items():
            text = text.replace(old, new)

        return text

    def _markdown_to_latex(self, content: str) -> str:
        """Convert markdown content to LaTeX"""
        # Convert headings
        content = re.sub(r'^### (.+)$', r'\\subsubsection{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'\\subsection{\1}', content, flags=re.MULTILINE)
        content = re.sub(r'^# (.+)$', r'\\section{\1}', content, flags=re.MULTILINE)

        # Convert bold
        content = re.sub(r'\*\*(.+?)\*\*', r'\\textbf{\1}', content)

        # Convert italic
        content = re.sub(r'\*(.+?)\*', r'\\textit{\1}', content)

        # Convert bullet points
        content = re.sub(r'^\- (.+)$', r'\\item \1', content, flags=re.MULTILINE)

        return content

    def export_to_markdown(self, proposal_data: Dict[str, Any]) -> str:
        """
        Export proposal to clean markdown
        """
        md_parts = []

        # Title
        md_parts.append(f"# {proposal_data.get('title', 'Grant Proposal')}")
        md_parts.append('')

        # Metadata
        md_parts.append(f"**Proposal Type:** {proposal_data.get('proposal_type', '')}")
        md_parts.append(f"**Keywords:** {proposal_data.get('keywords', '')}")
        md_parts.append(f"**Date:** {datetime.now().strftime('%B %d, %Y')}")
        md_parts.append('')
        md_parts.append('---')
        md_parts.append('')

        # Sections
        sections = proposal_data.get('sections', [])
        for section in sections:
            md_parts.append(f"## {section['name']}")
            md_parts.append('')
            md_parts.append(section.get('content', ''))
            md_parts.append('')
            md_parts.append('---')
            md_parts.append('')

        # Bibliography
        md_parts.append('## References')
        md_parts.append('')
        citations = proposal_data.get('bibliography', {}).get('apa', [])
        for idx, citation in enumerate(citations, 1):
            md_parts.append(f"{idx}. {citation}")
        md_parts.append('')

        return '\n'.join(md_parts)

    def export_to_html(self, proposal_data: Dict[str, Any]) -> str:
        """
        Export proposal to professional HTML
        """
        html_parts = []

        # HTML header
        html_parts.append('<!DOCTYPE html>')
        html_parts.append('<html lang="en">')
        html_parts.append('<head>')
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append(f'<title>{proposal_data.get("title", "Grant Proposal")}</title>')
        html_parts.append(self._get_html_styles())
        html_parts.append('</head>')
        html_parts.append('<body>')

        # Title page
        html_parts.append('<div class="cover-page">')
        html_parts.append(f'<h1>{proposal_data.get("title", "Grant Proposal")}</h1>')
        html_parts.append(f'<h2>{proposal_data.get("proposal_type", "")}</h2>')
        html_parts.append(f'<p class="metadata">Date: {datetime.now().strftime("%B %d, %Y")}</p>')
        html_parts.append('</div>')

        # Sections
        sections = proposal_data.get('sections', [])
        for section in sections:
            html_parts.append(f'<div class="section">')
            html_parts.append(f'<h2>{section["name"]}</h2>')

            content = section.get('content', '')
            # Convert markdown to HTML
            html_content = markdown.markdown(content)
            html_parts.append(html_content)

            html_parts.append('</div>')

        # Bibliography
        html_parts.append('<div class="section">')
        html_parts.append('<h2>References</h2>')
        html_parts.append('<ol class="references">')
        citations = proposal_data.get('bibliography', {}).get('apa', [])
        for citation in citations:
            html_parts.append(f'<li>{citation}</li>')
        html_parts.append('</ol>')
        html_parts.append('</div>')

        html_parts.append('</body>')
        html_parts.append('</html>')

        return '\n'.join(html_parts)

    def _get_html_styles(self) -> str:
        """Get CSS styles for HTML export"""
        return '''
<style>
    body {
        font-family: 'Times New Roman', Times, serif;
        max-width: 8.5in;
        margin: 0 auto;
        padding: 1in;
        line-height: 1.6;
    }
    .cover-page {
        text-align: center;
        padding: 3in 0;
    }
    .cover-page h1 {
        font-size: 24pt;
        color: #1f4e79;
        margin-bottom: 0.5in;
    }
    .cover-page h2 {
        font-size: 18pt;
        color: #4f81bd;
    }
    .metadata {
        font-size: 12pt;
        color: #666;
    }
    .section {
        page-break-inside: avoid;
        margin-bottom: 1.5em;
    }
    h2 {
        font-size: 16pt;
        color: #1f4e79;
        border-bottom: 2px solid #1f4e79;
        padding-bottom: 0.2em;
        margin-top: 1.5em;
    }
    p {
        text-align: justify;
        margin-bottom: 1em;
    }
    .references {
        padding-left: 1.5em;
    }
    .references li {
        margin-bottom: 0.5em;
    }
    @media print {
        body {
            padding: 0;
        }
        .section {
            page-break-inside: avoid;
        }
    }
</style>
'''
