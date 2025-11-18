from django.shortcuts import render, redirect
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import UserProfile, SavedProposal
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
import io

def register_view(request):
    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        email = request.POST.get('email', '').strip()
        first_name = request.POST.get('first_name', '').strip()
        last_name = request.POST.get('last_name', '').strip()
        university = request.POST.get('university', '').strip()
        password = request.POST.get('password', '').strip()
        confirm_password = request.POST.get('confirm_password', '').strip()
        
        # Validation
        if not all([username, email, first_name, last_name, university, password, confirm_password]):
            messages.error(request, 'All fields are required.')
            return render(request, 'register.html')
        
        if password != confirm_password:
            messages.error(request, 'Passwords do not match.')
            return render(request, 'register.html')
        
        if len(password) < 8:
            messages.error(request, 'Password must be at least 8 characters long.')
            return render(request, 'register.html')
        
        if User.objects.filter(username=username).exists():
            messages.error(request, 'Username already exists.')
            return render(request, 'register.html')
        
        if User.objects.filter(email=email).exists():
            messages.error(request, 'Email already registered.')
            return render(request, 'register.html')
        
        # Create user
        try:
            user = User.objects.create_user(
                username=username,
                email=email,
                password=password
            )
            
            # Update profile
            profile = user.profile
            profile.first_name = first_name
            profile.last_name = last_name
            profile.university = university
            profile.save()
            
            messages.success(request, 'Registration successful! Please login.')
            return redirect('login')
        
        except Exception as e:
            messages.error(request, f'Registration failed: {str(e)}')
            return render(request, 'register.html')
    
    return render(request, 'register.html')


def login_view(request):
    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()
        
        if not username or not password:
            messages.error(request, 'Please provide both username and password.')
            return render(request, 'login.html')
        
        user = authenticate(request, username=username, password=password)
        
        if user is not None:
            login(request, user)
            messages.success(request, f'Welcome back, {user.profile.first_name}!')
            next_url = request.GET.get('next', 'home')
            return redirect(next_url)
        else:
            messages.error(request, 'Invalid username or password.')
            return render(request, 'login.html')
    
    return render(request, 'login.html')


@login_required(login_url='login')
def logout_view(request):
    logout(request)
    messages.success(request, 'Logged out successfully.')
    return redirect('login')


@login_required(login_url='login')
def profile_view(request):
    user = request.user
    profile = user.profile
    proposals = SavedProposal.objects.filter(user=user)
    
    context = {
        'user': user,
        'profile': profile,
        'proposals': proposals,
        'total_proposals': proposals.count(),
        'total_words': sum(p.word_count for p in proposals),
    }
    
    return render(request, 'profile.html', context)


@login_required(login_url='login')
def save_proposal_view(request):
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()
        keywords = request.POST.get('keywords', '').strip()
        description = request.POST.get('description', '').strip()
        content = request.POST.get('content', '').strip()
        
        if not all([title, description, content]):
            messages.error(request, 'Title, description, and content are required.')
            return redirect('home')
        
        word_count = len(content.split())
        citation_count = content.count('[')
        
        SavedProposal.objects.create(
            user=request.user,
            title=title,
            keywords=keywords,
            description=description,
            content=content,
            word_count=word_count,
            citation_count=citation_count
        )
        
        messages.success(request, 'Proposal saved successfully!')
        return redirect('profile')
    
    return redirect('home')


@login_required(login_url='login')
def delete_proposal_view(request, proposal_id):
    try:
        proposal = SavedProposal.objects.get(id=proposal_id, user=request.user)
        proposal.delete()
        messages.success(request, 'Proposal deleted successfully.')
    except SavedProposal.DoesNotExist:
        messages.error(request, 'Proposal not found.')
    
    return redirect('profile')


@login_required(login_url='login')
def download_proposal_view(request, proposal_id):
    try:
        proposal = SavedProposal.objects.get(id=proposal_id, user=request.user)
        
        # Create DOCX
        doc = Document()
        
        # Add title
        title = doc.add_heading(proposal.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add keywords
        if proposal.keywords:
            p = doc.add_paragraph()
            p.add_run('Keywords: ').bold = True
            p.add_run(proposal.keywords)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content (clean HTML/formatting)
        content_text = strip_html_tags(proposal.content)
        doc.add_paragraph(content_text)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Return as download
        response = HttpResponse(
            buffer.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f"{proposal.title[:50].replace(' ', '_')}_Proposal.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except SavedProposal.DoesNotExist:
        messages.error(request, 'Proposal not found.')
        return redirect('profile')

def strip_html_tags(text):
    import re
    clean = re.compile('<.*?>')
    return re.sub(clean, '', text)